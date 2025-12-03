from pathlib import Path
from typing import Optional, Tuple
import tempfile
import os

def extract_text_from_file(file_path: Path, file_extension: str) -> Tuple[Optional[str], Optional[str]]:
    """
    Извлечь текст из файла в зависимости от расширения.

    Args:
        file_path: Путь к файлу
        file_extension: Расширение файла (.doc, .docx, .pdf)

    Returns:
        Кортеж (текст, ошибка). Если успешно - (текст, None), если ошибка - (None, текст_ошибки)
    """
    try:
        if file_extension == ".docx":
            return extract_from_docx(file_path)
        elif file_extension == ".doc":
            return extract_from_doc(file_path)
        elif file_extension == ".pdf":
            return extract_from_pdf(file_path)
        else:
            return None, "Неподдерживаемый формат файла. Допустимы: .doc, .docx, .pdf"
    except Exception as e:
        return None, f"Произошла ошибка при обработке файла: {str(e)}"

def extract_from_docx(file_path: Path) -> Tuple[Optional[str], Optional[str]]:
    """
    Извлечь текст из .docx файла.

    Args:
        file_path: Путь к файлу

    Returns:
        Кортеж (текст, ошибка)
    """
    try:
        from docx import Document

        doc = Document(str(file_path))

        # Извлечь текст из параграфов
        paragraphs_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text)

        # Извлечь текст из таблиц
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    tables_text.append(" | ".join(row_text))

        # Объединить весь текст
        full_text = "\n".join(paragraphs_text)
        if tables_text:
            full_text += "\n\n=== ТАБЛИЦЫ ===\n\n" + "\n".join(tables_text)

        if not full_text.strip():
            return None, "Файл не содержит текста или пуст."

        return full_text, None

    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            return None, "Файл защищен паролем. Снимите защиту и попробуйте снова."
        elif "corrupt" in error_msg or "invalid" in error_msg:
            return None, "Файл поврежден и не может быть обработан. Попробуйте другой файл."
        else:
            return None, f"Ошибка при чтении файла: {str(e)}"

def extract_from_doc(file_path: Path) -> Tuple[Optional[str], Optional[str]]:
    """
    Извлечь текст из .doc файла через COM (pywin32).

    ВАЖНО: Работает только на Windows с установленным Microsoft Word!

    Args:
        file_path: Путь к файлу

    Returns:
        Кортеж (текст, ошибка)
    """
    try:
        import pythoncom
        import win32com.client

        # Инициализация COM для текущего потока
        pythoncom.CoInitialize()

        try:
            # Создать экземпляр Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False  # Не показывать окно Word

            # Открыть документ
            # Преобразуем путь в абсолютный для COM
            abs_path = str(file_path.resolve())

            try:
                doc = word.Documents.Open(abs_path, ReadOnly=True, PasswordDocument="")
            except Exception as e:
                if "password" in str(e).lower():
                    return None, "Файл защищен паролем. Снимите защиту и попробуйте снова."
                raise

            try:
                # Извлечь весь текст
                full_text = doc.Content.Text

                # Извлечь текст из таблиц (они могут не попасть в Content.Text)
                tables_text = []
                for table in doc.Tables:
                    table_content = []
                    for row in table.Rows:
                        row_text = []
                        for cell in row.Cells:
                            cell_text = cell.Range.Text.strip()
                            # Убрать служебные символы Word
                            cell_text = cell_text.replace('\r', '').replace('\x07', '')
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            table_content.append(" | ".join(row_text))
                    if table_content:
                        tables_text.append("\n".join(table_content))

                # Объединить текст
                result_text = full_text.strip()
                if tables_text:
                    result_text += "\n\n=== ТАБЛИЦЫ ===\n\n" + "\n\n".join(tables_text)

                if not result_text.strip():
                    return None, "Файл не содержит текста или пуст."

                return result_text, None

            finally:
                # Закрыть документ без сохранения
                doc.Close(SaveChanges=False)

        finally:
            # Закрыть Word
            word.Quit()
            pythoncom.CoUninitialize()

    except ImportError:
        return None, "Модуль pywin32 не установлен. Установите: pip install pywin32"
    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg:
            return None, "Файл защищен паролем. Снимите защиту и попробуйте снова."
        elif "corrupt" in error_msg or "cannot open" in error_msg:
            return None, "Файл поврежден и не может быть обработан. Попробуйте другой файл."
        elif "word" in error_msg and ("not found" in error_msg or "failed" in error_msg):
            return None, "Microsoft Word не установлен на сервере. Используйте .docx или .pdf формат."
        else:
            return None, f"Ошибка при чтении .doc файла: {str(e)}"

def extract_from_pdf(file_path: Path) -> Tuple[Optional[str], Optional[str]]:
    """
    Извлечь текст из .pdf файла.

    Args:
        file_path: Путь к файлу

    Returns:
        Кортеж (текст, ошибка)
    """
    try:
        import PyPDF2

        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)

            # Проверка на защиту паролем
            if pdf_reader.is_encrypted:
                return None, "Файл защищен паролем. Снимите защиту и попробуйте снова."

            # Извлечь текст со всех страниц
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, start=1):
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(f"=== Страница {page_num} ===\n{page_text}")

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                return None, "PDF не содержит текстового слоя (возможно, это скан). Используйте .docx"

            return full_text, None

    except Exception as e:
        error_msg = str(e).lower()
        if "password" in error_msg or "encrypted" in error_msg:
            return None, "Файл защищен паролем. Снимите защиту и попробуйте снова."
        elif "corrupt" in error_msg or "invalid" in error_msg or "damaged" in error_msg:
            return None, "Файл поврежден и не может быть обработан. Попробуйте другой файл."
        else:
            return None, f"Ошибка при чтении PDF: {str(e)}"

def estimate_token_count(text: str) -> int:
    """
    Оценить количество токенов в тексте.

    Используем простую эвристику: ~1 токен = 4 символа для русского текста.

    Args:
        text: Текст для оценки

    Returns:
        Примерное количество токенов
    """
    return len(text) // 4

def check_text_size(text: str, max_tokens: int = 128000) -> Tuple[bool, Optional[str]]:
    """
    Проверить размер текста на превышение лимита токенов.

    Args:
        text: Текст для проверки
        max_tokens: Максимальное количество токенов (по умолчанию 128K)

    Returns:
        Кортеж (результат, ошибка). True если размер допустим, False если превышен
    """
    token_count = estimate_token_count(text)

    if token_count > max_tokens:
        pages_estimate = token_count // 1500  # Примерно 1500 токенов = 1 страница
        return False, f"Документ слишком большой для обработки (примерно {pages_estimate} страниц, лимит: 100 страниц)."

    return True, None


def create_word_document(content: str, filename: str = "Анализ_договора", content_type: str = "markdown") -> "BytesIO":
    """
    Создать Word документ из контента с сохранением структуры.

    Args:
        content: Содержимое (Markdown или HTML)
        filename: Имя файла (без расширения)
        content_type: Тип контента ('markdown' или 'html')

    Returns:
        BytesIO объект с содержимым .docx файла
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from io import BytesIO
    from bs4 import BeautifulSoup, NavigableString, Tag
    import re
    import tempfile
    import os

    # Очистить filename от проблемных символов
    filename = re.sub(r'[^\w\-_\.]', '_', filename, flags=re.UNICODE)

    doc = Document()

    # Установить стили по умолчанию для документа
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    def set_cell_shading(cell, color: str):
        """Установить цвет фона ячейки таблицы."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)

    def add_formatted_runs(element, paragraph):
        """
        Рекурсивно добавляет текст с форматированием в параграф.
        Обрабатывает вложенные теги форматирования.
        """
        for child in element.children:
            if isinstance(child, NavigableString):
                text = str(child)
                if text.strip() or text == ' ':
                    # Сохраняем пробелы между словами
                    run = paragraph.add_run(text)
            elif isinstance(child, Tag):
                if child.name in ['strong', 'b']:
                    # Жирный текст
                    run = paragraph.add_run(child.get_text())
                    run.bold = True
                elif child.name in ['em', 'i']:
                    # Курсив
                    run = paragraph.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'code':
                    # Код - моноширинный шрифт
                    run = paragraph.add_run(child.get_text())
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif child.name == 'a':
                    # Ссылка - подчеркнутый синий текст
                    run = paragraph.add_run(child.get_text())
                    run.underline = True
                    run.font.color.rgb = RGBColor(0, 0, 255)
                elif child.name == 'br':
                    # Перенос строки
                    paragraph.add_run('\n')
                elif child.name in ['span', 'u', 's', 'del', 'ins', 'mark', 'sub', 'sup']:
                    # Прочие inline элементы - рекурсивно обрабатываем
                    add_formatted_runs(child, paragraph)
                else:
                    # Для других тегов просто извлекаем текст
                    text = child.get_text()
                    if text.strip():
                        paragraph.add_run(text)

    def process_table(table_element, doc):
        """Создаёт настоящую таблицу Word из HTML таблицы."""
        rows = table_element.find_all('tr')
        if not rows:
            return

        # Определяем количество колонок
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            if len(cells) > max_cols:
                max_cols = len(cells)

        if max_cols == 0:
            return

        # Создаём таблицу Word
        word_table = doc.add_table(rows=len(rows), cols=max_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заполняем таблицу
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for col_idx, cell in enumerate(cells):
                if col_idx < max_cols:
                    word_cell = word_table.rows[row_idx].cells[col_idx]

                    # Очищаем дефолтный параграф
                    word_cell.paragraphs[0].clear()

                    # Добавляем форматированный текст
                    add_formatted_runs(cell, word_cell.paragraphs[0])

                    # Если это заголовок таблицы (th), делаем жирным и добавляем фон
                    if cell.name == 'th':
                        for run in word_cell.paragraphs[0].runs:
                            run.bold = True
                        set_cell_shading(word_cell, 'D9E2F3')  # Светло-синий фон

        # Добавляем пустую строку после таблицы
        doc.add_paragraph()

    def process_list(list_element, doc, is_ordered=False, level=0):
        """Обрабатывает списки ul/ol с поддержкой вложенности."""
        items = list_element.find_all('li', recursive=False)

        for idx, item in enumerate(items, 1):
            # Формируем отступ для вложенных списков
            indent = "    " * level

            # Определяем маркер
            if is_ordered:
                marker = f"{idx}."
            else:
                markers = ['•', '◦', '▪', '▹']
                marker = markers[min(level, len(markers) - 1)]

            # Создаём параграф для элемента списка
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25 * (level + 1))

            # Добавляем маркер
            run = para.add_run(f"{marker} ")
            if is_ordered:
                run.bold = True

            # Добавляем содержимое элемента (без вложенных списков)
            for child in item.children:
                if isinstance(child, NavigableString):
                    text = str(child).strip()
                    if text:
                        para.add_run(text)
                elif isinstance(child, Tag):
                    if child.name in ['ul', 'ol']:
                        # Вложенный список обработаем отдельно
                        continue
                    elif child.name in ['strong', 'b']:
                        run = para.add_run(child.get_text())
                        run.bold = True
                    elif child.name in ['em', 'i']:
                        run = para.add_run(child.get_text())
                        run.italic = True
                    elif child.name == 'code':
                        run = para.add_run(child.get_text())
                        run.font.name = 'Courier New'
                    else:
                        text = child.get_text()
                        if text.strip():
                            para.add_run(text)

            # Обрабатываем вложенные списки
            for child in item.children:
                if isinstance(child, Tag):
                    if child.name == 'ul':
                        process_list(child, doc, is_ordered=False, level=level + 1)
                    elif child.name == 'ol':
                        process_list(child, doc, is_ordered=True, level=level + 1)

    def process_element(element, doc, processed_elements):
        """
        Рекурсивно обрабатывает элементы HTML в порядке их появления.
        """
        if element in processed_elements:
            return
        processed_elements.add(element)

        if isinstance(element, NavigableString):
            return

        if not isinstance(element, Tag):
            return

        tag_name = element.name

        # Заголовки
        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag_name[1])
            level = min(level, 4)  # Word поддерживает уровни 0-4
            heading = doc.add_heading(level=level)
            add_formatted_runs(element, heading)

        # Параграфы
        elif tag_name == 'p':
            para = doc.add_paragraph()
            add_formatted_runs(element, para)

        # Таблицы
        elif tag_name == 'table':
            process_table(element, doc)

        # Ненумерованные списки
        elif tag_name == 'ul':
            process_list(element, doc, is_ordered=False)

        # Нумерованные списки
        elif tag_name == 'ol':
            process_list(element, doc, is_ordered=True)

        # Цитаты
        elif tag_name == 'blockquote':
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            # Добавляем вертикальную линию через форматирование
            for run in para.runs:
                run.italic = True
            add_formatted_runs(element, para)
            for run in para.runs:
                run.italic = True

        # Горизонтальная линия
        elif tag_name == 'hr':
            para = doc.add_paragraph('─' * 50)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Блоки кода
        elif tag_name == 'pre':
            code_element = element.find('code')
            if code_element:
                code_text = code_element.get_text()
            else:
                code_text = element.get_text()

            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

        # div и другие контейнеры - обрабатываем содержимое
        elif tag_name in ['div', 'section', 'article', 'main', 'body', 'html']:
            for child in element.children:
                if isinstance(child, Tag):
                    process_element(child, doc, processed_elements)

        # Переносы строк как отдельный параграф
        elif tag_name == 'br':
            doc.add_paragraph()

    # ===== ОСНОВНАЯ ЛОГИКА =====

    if content_type == "markdown":
        # Конвертируем Markdown в HTML
        import markdown

        try:
            # Используем расширения для таблиц
            html_content = markdown.markdown(
                content,
                extensions=['tables', 'fenced_code', 'nl2br']
            )
            content = html_content
            content_type = "html"
        except ImportError:
            # Fallback если markdown не установлен - простая обработка
            lines = content.split('\n')
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    doc.add_paragraph()
                elif line_stripped.startswith('#'):
                    level = len(re.match(r'^#+', line_stripped).group())
                    title = line_stripped.lstrip('#').strip()
                    doc.add_heading(title, level=min(level, 4))
                elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                    item = line_stripped.lstrip('- ').lstrip('* ')
                    doc.add_paragraph(f"• {item}")
                elif re.match(r'^\d+\.', line_stripped):
                    doc.add_paragraph(line_stripped)
                else:
                    doc.add_paragraph(line_stripped)

            # Сохраняем и возвращаем
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
            try:
                os.close(tmp_fd)
                doc.save(tmp_path)
                with open(tmp_path, 'rb') as f:
                    doc_io = BytesIO(f.read())
                    doc_io.seek(0)
                return doc_io
            finally:
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

    if content_type == "html":
        try:
            # Парсим HTML
            if not content.strip().startswith('<'):
                content = f"<div>{content}</div>"

            soup = BeautifulSoup(content, 'html.parser')

            # Множество для отслеживания обработанных элементов
            processed_elements = set()

            # Находим корневой элемент
            root = soup.body if soup.body else soup

            # Обрабатываем все элементы верхнего уровня в порядке появления
            for element in root.children:
                if isinstance(element, Tag):
                    process_element(element, doc, processed_elements)

        except ImportError as e:
            # BeautifulSoup не установлен
            doc.add_paragraph("Ошибка: BeautifulSoup не установлен.")
            doc.add_paragraph(re.sub(r'<[^>]+>', '', content))
        except Exception as e:
            # Fallback при ошибке - извлекаем просто текст
            print(f"Ошибка обработки HTML: {e}")
            try:
                soup = BeautifulSoup(content, 'html.parser')
                text = soup.get_text(separator='\n')
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            except Exception:
                doc.add_paragraph(re.sub(r'<[^>]+>', '', content))

    # Сохраняем документ через временный файл
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    try:
        os.close(tmp_fd)
        doc.save(tmp_path)

        with open(tmp_path, 'rb') as f:
            doc_io = BytesIO(f.read())
            doc_io.seek(0)

        return doc_io
    finally:
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except OSError:
            pass